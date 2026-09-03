"""End-to-end ingest. Owner: person 2.

Builds real PDFs with the corpus writer and runs them through the whole
pipeline, because the interesting failures in ingest are all at the seams:
a native page that is wrongly classified as a scan, a table whose rows are
found but whose cells are not, a page number that is off by one.

The scanned case is marked `slow` -- it loads an OCR model and runs it -- so
`pytest -m "not slow"` stays fast for the inner loop while the full run still
exercises the path the demo actually depends on.
"""

from __future__ import annotations

import sys
from pathlib import Path

import pytest

RAGSVC = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(RAGSVC / "eval"))

from pdfgen import PdfWriter, scan_effect  # noqa: E402

from ingest.pipeline import ingest_document  # noqa: E402


def sample_document() -> PdfWriter:
    writer = PdfWriter(title="Thickness Survey", doc_no="TMS/2024/TEST/01")
    writer.add_title("PIPING THICKNESS SURVEY", "Test circuit")
    writer.add_key_values([("Survey No.", "TMS/2024/TEST/01"), ("Line", "8-P-1104")])
    writer.add_heading("1. Circuit Description")
    writer.add_paragraph(
        "The retirement thickness for line 8-P-1104 is 5.2 mm, derived from the "
        "pressure design thickness plus a structural margin."
    )
    writer.add_heading("2. Measurement Results")
    writer.add_table(
        ["CML", "Component", "Nominal", "Measured", "Rate"],
        [
            ["CML-01", "Straight run A", "8.18", "7.91", "0.04"],
            ["CML-12", "Straight run D", "8.18", "7.80", "0.12"],
            ["CML-19", "Low point drain", "6.35", "5.62", "0.24"],
        ],
    )
    return writer


@pytest.fixture
def native_pdf(tmp_path: Path) -> Path:
    path = tmp_path / "native-survey.pdf"
    writer = sample_document()
    writer.document.save(str(path))
    writer.document.close()
    return path


@pytest.fixture
def scanned_pdf(tmp_path: Path) -> Path:
    path = tmp_path / "scanned-survey.pdf"
    writer = sample_document()
    path.write_bytes(scan_effect(writer.to_bytes(), seed=1))
    writer.document.close()
    return path


# --- native path ------------------------------------------------------------


def test_native_pdf_skips_ocr_entirely(native_pdf: Path):
    result = ingest_document(native_pdf)

    assert result.page_count >= 1
    assert result.native_pages == result.page_count
    assert result.scanned_pages == 0
    assert result.ocr_backend == "none", "OCR ran on a page that already had text"


def test_native_ingest_recovers_the_table_as_markdown(native_pdf: Path):
    result = ingest_document(native_pdf)
    text = "\n".join(page.text for page in result.pages)

    assert "CML-19" in text and "5.62" in text
    assert "|" in text, "the table was flattened instead of kept as a table"


def test_chunks_come_back_with_provenance(native_pdf: Path):
    result = ingest_document(native_pdf, doc_id="doc-test", filename="native-survey.pdf")

    assert result.chunks
    for chunk in result.chunks:
        assert chunk.doc_id == "doc-test"
        assert chunk.filename == "native-survey.pdf"
        assert 1 <= chunk.page <= result.page_count


def test_the_table_lands_in_exactly_one_chunk(native_pdf: Path):
    result = ingest_document(native_pdf)
    holders = [c for c in result.chunks if "CML-19" in c.text]
    assert len(holders) == 1


def test_digest_and_size_are_recorded(native_pdf: Path):
    result = ingest_document(native_pdf)
    assert len(result.sha256) == 64
    assert result.size_bytes == native_pdf.stat().st_size


# --- scanned path -----------------------------------------------------------


@pytest.mark.slow
def test_scanned_pdf_is_routed_through_ocr(scanned_pdf: Path):
    result = ingest_document(scanned_pdf)

    assert result.scanned_pages == result.page_count
    assert result.native_pages == 0
    if result.ocr_error:
        pytest.skip(f"no OCR backend installed: {result.ocr_error}")
    assert result.ocr_backend in {"paddle", "rapidocr"}


@pytest.mark.slow
def test_ocr_recovers_the_identifiers_that_queries_use(scanned_pdf: Path):
    result = ingest_document(scanned_pdf)
    if result.ocr_error:
        pytest.skip("no OCR backend installed")

    text = "\n".join(page.text for page in result.pages)
    # Tag numbers are what refinery staff actually search for. If OCR loses
    # them, BM25 has nothing to match and hybrid retrieval buys nothing.
    assert "CML" in text
    assert "8-P-1104" in text.replace(" ", "") or "8-P-1104" in text


@pytest.mark.slow
def test_scanned_pages_carry_a_confidence_score(scanned_pdf: Path):
    result = ingest_document(scanned_pdf)
    if result.ocr_error:
        pytest.skip("no OCR backend installed")
    for page in result.pages:
        assert 0.0 <= page.mean_conf <= 1.0
    assert result.mean_conf > 0.5, "OCR confidence collapsed on a legible page"


@pytest.mark.slow
def test_render_dpi_is_recorded_for_scanned_pages(scanned_pdf: Path):
    result = ingest_document(scanned_pdf)
    assert result.dpi_used
    assert all(dpi in (150, 200) for dpi in result.dpi_used)


# --- images -----------------------------------------------------------------


def test_an_image_file_is_accepted_as_a_one_page_document(tmp_path: Path):
    import numpy as np
    import cv2

    image = np.full((600, 800, 3), 255, dtype=np.uint8)
    cv2.putText(image, "PSV-2103", (60, 300), cv2.FONT_HERSHEY_SIMPLEX, 2.0, (0, 0, 0), 3)
    path = tmp_path / "photo.png"
    cv2.imwrite(str(path), image)

    result = ingest_document(path)
    assert result.page_count == 1


# --- text-native formats (ingest/textdoc.py) ----------------------------------

def test_textdoc_big_sheet_is_split_by_rows_with_header_repeated(tmp_path):
    """A 3,000-row workbook must not become one atomic table block.

    Table blocks are never split downstream, so an unsplit sheet is a single
    block larger than the whole tool budget: unretrievable as one hit and
    truncated to nothing as a read. Each piece carries the header so a hit
    is self-describing.
    """
    import openpyxl
    from ingest.textdoc import TABLE_ROWS_PER_BLOCK, pages_from_text_document

    wb = openpyxl.Workbook(); ws = wb.active; ws.title = "refcap"
    ws.append(["SITE", "STATE", "CAPACITY_BPD"])
    for i in range(3000):
        ws.append([f"SITE-{i}", "Texas", 1000 + i])
    path = tmp_path / "big.xlsx"; wb.save(path)

    pages = pages_from_text_document(path)
    tables = [b for p in pages for b in p.blocks if b.kind == "table"]
    assert len(tables) == -(-3000 // TABLE_ROWS_PER_BLOCK)
    assert all("| SITE | STATE | CAPACITY_BPD |" in b.text for b in tables)
    assert "SITE-2999" in tables[-1].text
    assert "(rows 1-40 of 3000)" in tables[0].text
    assert len(pages) > 1


def test_textdoc_docx_headings_paragraphs_tables(tmp_path):
    import docx
    from ingest.textdoc import pages_from_text_document

    d = docx.Document(); d.add_heading("Valve register", 1); d.add_paragraph("Debutaniser section.")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Tag"; t.cell(0, 1).text = "Set pressure"
    t.cell(1, 0).text = "PSV-2103"; t.cell(1, 1).text = "12.5 barg"
    path = tmp_path / "reg.docx"; d.save(path)

    pages = pages_from_text_document(path)
    assert [b.kind for b in pages[0].blocks] == ["heading", "paragraph", "table"]
    assert "| PSV-2103 | 12.5 barg |" in pages[0].blocks[2].text
    assert pages[0].blocks[2].section == "Valve register"
