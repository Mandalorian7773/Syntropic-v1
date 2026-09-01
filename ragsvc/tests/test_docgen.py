"""Deliverable generation. Owner: person 2.

The claim being tested is the one in the problem statement: the output is a
real file. So these tests open the generated files back up with the same
libraries Word and Excel-compatible readers use, and check the structure a
reviewer would look for -- a reference number, the required headings, a
signature block, numbers stored as numbers.
"""

from __future__ import annotations

from docx import Document
from openpyxl import load_workbook

import ragconfig as cfg
import ragdb
from docgen import Section, Sheet, available_templates, create_docx, create_xlsx, load_template


def docx_text(path: str) -> str:
    document = Document(path)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    for section in document.sections:
        parts.extend(p.text for p in section.footer.paragraphs)
    return "\n".join(parts)


# --- templates --------------------------------------------------------------


def test_the_three_required_templates_exist():
    assert {"approval_note", "inspection_summary", "calculation_sheet"} <= set(
        available_templates()
    )


def test_every_template_declares_a_signature_block_and_required_sections():
    for name in available_templates():
        template = load_template(name)
        assert template.signature_block.get("columns"), f"{name} has no signature block"
        assert any(s.required for s in template.sections), f"{name} has no required section"
        assert template.reference_prefix


# --- approval note ----------------------------------------------------------


def test_approval_note_contains_the_furniture_of_a_real_note():
    record = create_docx(
        "approval_note",
        "Repair of Nozzle N1 on Vessel V-1201",
        [
            Section(key="background", body="Corrosion under insulation was confirmed at N1."),
            Section(key="observations", body="Minimum remaining thickness 8.6 mm against 9.4 mm required."),
            Section(key="recommendation", body="Replace the affected nozzle neck during the February shutdown."),
        ],
        meta={"originator": "S. Rajagopal", "priority": "High"},
    )
    text = docx_text(record.path)

    assert cfg.ORG_NAME.upper() in text
    assert "Repair of Nozzle N1 on Vessel V-1201" in text
    assert "MRPL/I&R/APR/" in text, "no reference number in the house format"
    assert "1. Background" in text
    assert "4. Recommendation" in text
    assert "Prepared By" in text and "Reviewed By" in text and "Approved By" in text
    assert "Distribution" in text
    assert cfg.GENERATED_NOTICE in text, "a generated file must say that it is one"


def test_required_sections_are_printed_even_when_the_caller_omits_them():
    """A note missing its Recommendation must look incomplete, not complete."""
    record = create_docx(
        "approval_note", "Partial note", [Section(key="background", body="Only background.")]
    )
    text = docx_text(record.path)

    assert "4. Recommendation" in text
    assert "not stated" in text


def test_optional_sections_are_omitted_when_empty():
    record = create_docx(
        "approval_note", "Lean note", [Section(key="background", body="Background only.")]
    )
    text = docx_text(record.path)
    assert "5. Financial Implication" not in text


def test_unmatched_sections_from_the_model_are_kept_not_discarded():
    record = create_docx(
        "approval_note",
        "Note with an extra section",
        [
            Section(key="background", body="B."),
            Section(heading="Site Photographs", body="Three photographs are attached."),
        ],
    )
    text = docx_text(record.path)
    assert "Site Photographs" in text
    assert "Three photographs are attached." in text


# --- inspection summary -----------------------------------------------------


def test_inspection_summary_renders_the_findings_table():
    record = create_docx(
        "inspection_summary",
        "External inspection of V-1201",
        [
            Section(key="scope", body="Close visual examination and UT thickness survey."),
            Section(
                key="findings",
                table=[
                    ["1", "Damaged insulation cladding at N1", "Major", "INS/2024/0117 p.4", "Strip and inspect", "31 Mar 2024"],
                    ["2", "Surface rust on skirt", "Minor", "INS/2024/0117 p.4", "Touch up", "30 Jun 2024"],
                ],
            ),
            Section(key="recommendation", body="Increase inspection frequency at CML-07 to annual."),
        ],
        meta={"equipment": "V-1201", "unit": "CDU-I"},
    )
    text = docx_text(record.path)

    assert "Damaged insulation cladding at N1" in text
    assert "Severity" in text and "Major" in text
    assert "V-1201" in text
    assert "MRPL/I&R/INS/" in text


def test_a_repeated_header_row_from_the_model_is_not_rendered_twice():
    record = create_docx(
        "inspection_summary",
        "Header duplication",
        [
            Section(
                key="findings",
                table=[
                    ["S.No", "Observation", "Severity", "Source Reference", "Recommended Action", "Target Date"],
                    ["1", "Only real finding", "Minor", "ref", "action", "date"],
                ],
            )
        ],
    )
    text = docx_text(record.path)
    assert text.count("Observation") == 1


# --- calculation sheet ------------------------------------------------------


def test_calculation_sheet_shows_formula_substitution_and_result():
    record = create_docx(
        "calculation_sheet",
        "Minimum required thickness of shell course 3, V-1201",
        [
            Section(
                key="inputs",
                table=[
                    ["P", "Design pressure", "10.5", "barg", "DS-V-1201 Rev.3"],
                    ["R", "Inside radius", "900", "mm", "Fabrication drawing"],
                    ["S", "Allowable stress", "138", "MPa", "ASME II-D"],
                    ["E", "Joint efficiency", "1.00", "-", "100% radiography"],
                ],
            ),
            Section(key="formula", body="t = (P x R) / (S x E - 0.6 x P)"),
            Section(key="substitution", body="t = (1.05 x 900) / (138 x 1.00 - 0.6 x 1.05)"),
            Section(key="result", body="t = 6.88 mm required; measured 11.4 mm; acceptable"),
            Section(
                key="assumptions",
                bullets=[
                    "Corrosion allowance of 3.0 mm is additional to the calculated thickness.",
                    "No external loading other than static head is considered.",
                ],
            ),
        ],
        meta={"equipment": "V-1201", "standard": "ASME VIII Div.1 UG-27"},
    )
    text = docx_text(record.path)

    assert "2. Formula" in text
    assert "t = (P x R) / (S x E - 0.6 x P)" in text
    assert "3. Substitution" in text
    assert "6.88 mm" in text
    assert "6. Assumptions" in text
    assert "Corrosion allowance of 3.0 mm" in text


# --- workbooks --------------------------------------------------------------


def test_xlsx_stores_numbers_as_numbers():
    record = create_xlsx(
        [
            Sheet(
                name="Thickness",
                title="CDU-P-11 survey 2024",
                columns=["CML", "Measured", "Rate"],
                rows=[["CML-12", "7.80", "0.12"], ["CML-19", "5.62", "0.24"]],
                notes=["Source: TMS/2024/CDU/03"],
            )
        ]
    )
    workbook = load_workbook(record.path)
    sheet = workbook["Thickness"]

    header_row = next(
        r for r in range(1, 10) if sheet.cell(row=r, column=1).value == "CML"
    )
    assert sheet.cell(row=header_row + 1, column=1).value == "CML-12"
    assert sheet.cell(row=header_row + 1, column=2).value == 7.80
    assert isinstance(sheet.cell(row=header_row + 2, column=3).value, float)
    assert sheet.freeze_panes is not None
    assert sheet.auto_filter.ref


def test_xlsx_handles_awkward_sheet_names():
    record = create_xlsx(
        [
            Sheet(name="Findings/2024: CDU [main]", columns=["A"], rows=[["1"]]),
            Sheet(name="Findings/2024: CDU [main]", columns=["A"], rows=[["2"]]),
        ]
    )
    workbook = load_workbook(record.path)
    assert len(workbook.sheetnames) == 2
    for name in workbook.sheetnames:
        assert len(name) <= 31
        assert not set(name) & set(":\\/?*[]")


# --- storage ----------------------------------------------------------------


def test_artifacts_are_recorded_and_stored_one_directory_each():
    from pathlib import Path

    first = create_docx("approval_note", "Same name", [Section(key="background", body="a")])
    second = create_docx("approval_note", "Same name", [Section(key="background", body="b")])

    assert first.artifact_id != second.artifact_id
    assert first.filename == second.filename
    assert Path(first.path).exists() and Path(second.path).exists()
    assert Path(first.path).parent != Path(second.path).parent

    row = ragdb.get_artifact(first.artifact_id)
    assert row["mime"].endswith("wordprocessingml.document")
    assert row["size_bytes"] > 0


def test_reference_numbers_increment_within_the_year():
    first = create_docx("approval_note", "First", [Section(key="background", body="a")])
    second = create_docx("approval_note", "Second", [Section(key="background", body="b")])

    def reference(path: str) -> str:
        text = docx_text(path)
        return next(line for line in text.splitlines() if "MRPL/I&R/" in line)

    assert reference(first.path) != reference(second.path)
