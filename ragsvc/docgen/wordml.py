"""Low-level OOXML helpers for python-docx. Owner: person 2.

python-docx covers paragraphs, runs and tables and stops there. Everything that
makes a document look like real paperwork -- ruled table borders, shaded header
rows, a rule under the letterhead, "Page 2 of 6" in the footer -- is below its
API and has to be built as XML.

That is what this module is for, and why it is separate: docx_builder.py should
read as document structure, not as namespace declarations.

Every function here takes a python-docx object and mutates it in place.
"""

from __future__ import annotations

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BORDER_SIDES = ("top", "left", "bottom", "right", "insideH", "insideV")


def _border(tag: str, size: int, colour: str, style: str = "single"):
    element = OxmlElement(f"w:{tag}")
    element.set(qn("w:val"), style)
    element.set(qn("w:sz"), str(size))  # eighths of a point
    element.set(qn("w:space"), "0")
    element.set(qn("w:color"), colour)
    return element


def set_table_borders(table, size: int = 6, colour: str = "808080") -> None:
    """Rule every edge of a table, inside and out."""
    properties = table._element.tblPr
    existing = properties.find(qn("w:tblBorders"))
    if existing is not None:
        properties.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for side in BORDER_SIDES:
        borders.append(_border(side, size, colour))
    properties.append(borders)


def clear_table_borders(table) -> None:
    """Remove all rules. Used for the header field grid, which is a layout
    device rather than a table the reader is meant to see."""
    properties = table._element.tblPr
    existing = properties.find(qn("w:tblBorders"))
    if existing is not None:
        properties.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for side in BORDER_SIDES:
        borders.append(_border(side, 0, "auto", style="none"))
    properties.append(borders)


def shade_cell(cell, colour: str) -> None:
    """Fill a table cell with a solid colour, given as RRGGBB."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), colour)
    cell._tc.get_or_add_tcPr().append(shading)


def cell_vertical_center(cell) -> None:
    alignment = OxmlElement("w:vAlign")
    alignment.set(qn("w:val"), "center")
    cell._tc.get_or_add_tcPr().append(alignment)


def paragraph_rule(paragraph, position: str = "bottom", size: int = 12,
                   colour: str = "1F3864") -> None:
    """Draw a horizontal rule on one edge of a paragraph.

    This is how the line under a letterhead is done in Word -- there is no
    "horizontal rule" object, only paragraph borders.
    """
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    borders.append(_border(position, size, colour))


def keep_with_next(paragraph) -> None:
    """Stop a heading from being orphaned at the foot of a page."""
    properties = paragraph._p.get_or_add_pPr()
    element = OxmlElement("w:keepNext")
    properties.append(element)


def repeat_header_row(row) -> None:
    """Mark a table row as a header so it repeats across page breaks.

    A findings table that runs onto page 3 with no column headings is a table
    nobody can read, and inspection tables run long.
    """
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def add_page_number_field(paragraph, prefix: str = "Page ", middle: str = " of ") -> None:
    """Insert a live "Page X of Y" field into a paragraph.

    Field codes rather than literal text: the page count is not known until
    Word paginates, and a hard-coded "Page 1 of 4" on a document that grew to
    six pages is the kind of detail a reviewer notices immediately.
    """
    def field(instruction: str) -> None:
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = instruction
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run = paragraph.add_run()
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(end)

    paragraph.add_run(prefix)
    field(" PAGE ")
    paragraph.add_run(middle)
    field(" NUMPAGES ")


def set_column_widths(table, widths_cm: list[float]) -> None:
    """Fix column widths.

    Word ignores a width set only on the table, so it is set on every cell of
    every row, which is what python-docx's own documentation recommends.
    """
    from docx.shared import Cm  # noqa: PLC0415

    table.autofit = False
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            if index < len(widths_cm):
                cell.width = Cm(widths_cm[index])
