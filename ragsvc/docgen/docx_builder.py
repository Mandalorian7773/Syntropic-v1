"""Word deliverable generation. Owner: person 2.

The problem statement asks for real files, not chat replies, and this is where
that claim is either true or not. A generated approval note that a refinery
manager would accept without editing is worth more to an evaluator than a
better retrieval score, so the effort here goes into the things that make
paperwork look like paperwork: a letterhead with a rule under it, a reference
number in the house format, a ruled field grid, numbered sections in the order
the department expects, and a three-column signature block at the foot.

Two decisions worth defending:

**Required sections are printed even when empty.** An approval note with no
Recommendation renders the heading and an explicit "not stated" placeholder.
Dropping the heading would produce a document that looks complete and is not.

**Every document is stamped.** The footer carries a notice that the file is a
system-generated draft with no valid signature. This service can write anything
it is asked to write; nothing it writes should be mistakable for a document a
human has already approved.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

import ragconfig as cfg

from . import templates as tpl
from . import wordml
from .schema import Section

BODY_FONT = "Arial"
BODY_SIZE = Pt(10)
HEADER_FILL = "D9D9D9"
BANNER_FILL = "1F3864"
PLACEHOLDER = "— not stated —"
GREY = RGBColor(0x80, 0x80, 0x80)
DARK = RGBColor(0x1F, 0x38, 0x64)


# --- page furniture ---------------------------------------------------------


def _setup_page(document: Document, orientation: str) -> None:
    section = document.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)  # A4
    if orientation == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.8)

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.paragraph_format.space_after = Pt(4)


def _letterhead(document: Document) -> None:
    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run(cfg.ORG_NAME.upper())
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = DARK
    name.paragraph_format.space_after = Pt(0)

    unit = document.add_paragraph()
    unit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    unit_run = unit.add_run(cfg.ORG_UNIT)
    unit_run.font.size = Pt(10)
    unit_run.bold = True
    unit.paragraph_format.space_after = Pt(0)

    where = document.add_paragraph()
    where.alignment = WD_ALIGN_PARAGRAPH.CENTER
    where_run = where.add_run(cfg.ORG_LOCATION)
    where_run.font.size = Pt(8)
    where_run.italic = True
    where.paragraph_format.space_after = Pt(2)
    wordml.paragraph_rule(where, "bottom", size=12)


def _banner(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    wordml.shade_cell(cell, BANNER_FILL)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Letter-spacing by hand: Word's character spacing is a run property that
    # python-docx does not expose, and a banner that reads A P P R O V A L
    # N O T E is the convention these documents follow.
    run = paragraph.add_run(" ".join(text.upper()))
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def _field_grid(document: Document, fields: list[tuple[str, str]]) -> None:
    """Two label/value pairs per row, borderless."""
    if not fields:
        return
    rows = (len(fields) + 1) // 2
    table = document.add_table(rows=rows, cols=4)
    wordml.clear_table_borders(table)
    wordml.set_column_widths(table, [3.4, 5.2, 3.2, 5.2])

    for index, (label, value) in enumerate(fields):
        row, column = index // 2, (index % 2) * 2
        label_cell = table.cell(row, column)
        label_paragraph = label_cell.paragraphs[0]
        label_run = label_paragraph.add_run(f"{label}:")
        label_run.bold = True
        label_run.font.size = Pt(9)

        value_cell = table.cell(row, column + 1)
        value_paragraph = value_cell.paragraphs[0]
        value_run = value_paragraph.add_run(str(value) if value else PLACEHOLDER)
        value_run.font.size = Pt(9)
        if not value:
            value_run.font.color.rgb = GREY
            value_run.italic = True
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _subject(document: Document, label: str, title: str) -> None:
    paragraph = document.add_paragraph()
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(10.5)
    title_run = paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(10.5)
    wordml.paragraph_rule(paragraph, "bottom", size=6, colour="808080")
    paragraph.paragraph_format.space_after = Pt(8)


# --- section bodies ---------------------------------------------------------


def _heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    wordml.keep_with_next(paragraph)


def _placeholder(document: Document, hint: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(PLACEHOLDER + (f"  ({hint})" if hint else ""))
    run.italic = True
    run.font.color.rgb = GREY
    run.font.size = Pt(9)


def _prose(document: Document, text: str) -> None:
    for block in [b.strip() for b in text.split("\n\n") if b.strip()]:
        paragraph = document.add_paragraph(block)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _bullets(document: Document, items: list[str], numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        text = str(item).strip()
        if not text:
            continue
        try:
            document.add_paragraph(text, style=style)
        except KeyError:
            # A stripped-down base template may not define the list styles.
            document.add_paragraph(("• " if not numbered else "") + text)


def _table(
    document: Document,
    rows: list[list[str]],
    columns: list[str],
    severity_colours: dict[str, str],
) -> None:
    """Render a bordered table with a shaded, repeating header row."""
    if not rows and not columns:
        return
    body = [[str(c) if c is not None else "" for c in row] for row in rows]

    if columns:
        header, data = columns, body
        # The model often re-emits the header as the first data row. Dropping
        # the duplicate is friendlier than rendering the header twice.
        if data and [c.strip().lower() for c in data[0]] == [
            c.strip().lower() for c in columns
        ]:
            data = data[1:]
    else:
        header, data = body[0], body[1:]

    width = max([len(header)] + [len(r) for r in data]) if data else len(header)
    header = header + [""] * (width - len(header))
    data = [r + [""] * (width - len(r)) for r in data]

    table = document.add_table(rows=1 + len(data), cols=width)
    wordml.set_table_borders(table)

    severity_column = next(
        (i for i, name in enumerate(header) if name.strip().lower() == "severity"), None
    )

    for index, name in enumerate(header):
        cell = table.cell(0, index)
        wordml.shade_cell(cell, HEADER_FILL)
        wordml.cell_vertical_center(cell)
        run = cell.paragraphs[0].add_run(str(name))
        run.bold = True
        run.font.size = Pt(9)
    wordml.repeat_header_row(table.rows[0])

    for row_index, row in enumerate(data, start=1):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            run = cell.paragraphs[0].add_run(str(value))
            run.font.size = Pt(9)
            if column_index == severity_column:
                colour = severity_colours.get(str(value).strip().title())
                if colour:
                    wordml.shade_cell(cell, colour)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _formula(document: Document, text: str) -> None:
    for line in [ln for ln in text.splitlines() if ln.strip()]:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(line.strip())
        run.italic = True
        run.font.size = Pt(11)
        run.font.name = "Cambria Math"
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(4)


def _result(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    wordml.set_table_borders(table, size=12, colour="1F3864")
    cell = table.cell(0, 0)
    wordml.shade_cell(cell, "EAF1FB")
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text.strip())
    run.bold = True
    run.font.size = Pt(11)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _render_section(
    document: Document,
    spec: tpl.TemplateSection | None,
    section: Section | None,
    severity_colours: dict[str, str],
) -> None:
    heading = (spec.heading if spec else "") or (section.heading if section else "")
    kind = spec.kind if spec else ("table" if section and section.table else "prose")
    if heading:
        _heading(document, heading)

    if section is None or section.is_empty:
        _placeholder(document, spec.hint if spec else "")
        return

    if section.table:
        _table(document, section.table, spec.columns if spec else [], severity_colours)
        if section.body.strip():
            _prose(document, section.body)
        return

    if kind == "formula":
        _formula(document, section.body)
    elif kind == "result":
        _result(document, section.body)
    elif kind == "numbered":
        items = section.bullets or [
            ln.strip() for ln in section.body.splitlines() if ln.strip()
        ]
        _bullets(document, items, numbered=True)
    else:
        if section.body.strip():
            _prose(document, section.body)
        if section.bullets:
            _bullets(document, section.bullets)


# --- foot of the document ---------------------------------------------------


def _signature_block(document: Document, block: dict) -> None:
    columns = block.get("columns") or []
    if not columns:
        return
    document.add_paragraph().paragraph_format.space_after = Pt(6)

    table = document.add_table(rows=2, cols=len(columns))
    wordml.set_table_borders(table)
    for index, column in enumerate(columns):
        header = table.cell(0, index)
        wordml.shade_cell(header, HEADER_FILL)
        run = header.paragraphs[0].add_run(column.get("role", ""))
        run.bold = True
        run.font.size = Pt(9)

        body = table.cell(1, index)
        first = True
        for label in column.get("lines", []):
            paragraph = body.paragraphs[0] if first else body.add_paragraph()
            first = False
            # A signature block needs room to actually sign in; the trailing
            # rule is what a pen goes on.
            line_run = paragraph.add_run(f"{label}: " + "_" * 18)
            line_run.font.size = Pt(9)
            paragraph.paragraph_format.space_after = Pt(10)


def _distribution(document: Document, distribution: dict, override: list[str] | None) -> None:
    entries = override if override is not None else distribution.get("default") or []
    if not entries:
        return
    _heading(document, distribution.get("label", "Distribution"))
    for index, entry in enumerate(entries, start=1):
        paragraph = document.add_paragraph(f"{index}. {entry}")
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.size = Pt(9)


def _footer(document: Document, reference: str) -> None:
    footer = document.sections[0].footer
    notice = footer.paragraphs[0]
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice_run = notice.add_run(cfg.GENERATED_NOTICE)
    notice_run.bold = True
    notice_run.font.size = Pt(7.5)
    notice_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    wordml.paragraph_rule(notice, "top", size=6, colour="808080")

    line = footer.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    reference_run = line.add_run(f"{reference}    |    ")
    reference_run.font.size = Pt(7.5)
    wordml.add_page_number_field(line)
    for run in line.runs:
        run.font.size = Pt(7.5)
        run.font.color.rgb = GREY


# --- entry point ------------------------------------------------------------


def _match_sections(
    template: tpl.Template, sections: list[Section]
) -> tuple[dict[str, Section], list[Section]]:
    """Map caller sections onto template slots; return (matched, leftovers).

    A 7B model will sometimes send `key="findings"`, sometimes
    `heading="2. Observations and Findings"`, and sometimes a heading of its
    own invention. The first two are matched; the third is kept and appended
    rather than discarded, because content the model produced is content the
    user asked for.
    """
    by_key: dict[str, Section] = {}
    leftovers: list[Section] = []
    keys = {s.key for s in template.sections}

    def normalise(text: str) -> str:
        return "".join(c for c in text.lower() if c.isalnum())

    heading_index = {normalise(s.heading): s.key for s in template.sections}
    heading_index.update({normalise(s.key): s.key for s in template.sections})

    for section in sections:
        target = None
        if section.key and section.key in keys:
            target = section.key
        elif section.heading:
            target = heading_index.get(normalise(section.heading))
        if target and target not in by_key:
            by_key[target] = section
        else:
            leftovers.append(section)
    return by_key, leftovers


def build_docx(
    template_name: str,
    title: str,
    sections: list[Section],
    path: Path,
    meta: dict | None = None,
) -> str:
    """Write a .docx to `path`. Returns the reference number it was given."""
    meta = dict(meta or {})
    template = tpl.load(template_name)
    reference = meta.get("reference") or tpl.next_reference(
        template.reference_prefix, template.id
    )

    document = Document()
    _setup_page(document, template.orientation)
    _letterhead(document)
    _banner(document, template.document_type)

    fields: list[tuple[str, str]] = []
    for spec in template.header_fields:
        auto = spec.get("auto")
        if auto == "reference":
            value = reference
        elif auto == "date":
            value = meta.get(spec["key"]) or tpl.today()
        else:
            value = meta.get(spec["key"], spec.get("default", ""))
        fields.append((spec.get("label", spec["key"]), str(value or "")))
    _field_grid(document, fields)
    _subject(document, template.subject_label, title)

    matched, leftovers = _match_sections(template, sections)
    for spec in template.sections:
        section = matched.get(spec.key)
        if section is None and not spec.required:
            continue
        _render_section(document, spec, section, template.severity_colours)

    for section in leftovers:
        if not section.is_empty or section.heading:
            _render_section(document, None, section, template.severity_colours)

    _signature_block(document, template.signature_block)
    _distribution(document, template.distribution, meta.get("distribution"))
    _footer(document, reference)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    return reference
